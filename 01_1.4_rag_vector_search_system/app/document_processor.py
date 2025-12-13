"""
Phase 2: Document Processing and Embedding Pipeline
Handles document extraction, chunking, and embedding generation
"""

import boto3
import json
import os
import uuid
import hashlib
import io
import re
from typing import Dict, List, Tuple, Optional
from datetime import datetime
import logging

# Document processing libraries
try:
    import PyPDF2
    import docx
    from bs4 import BeautifulSoup
except ImportError:
    pass

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Processes documents and generates embeddings"""
    
    def __init__(self, region: str = 'us-east-1'):
        self.s3 = boto3.client('s3', region_name=region)
        self.region = region
    
    def extract_text_from_pdf(self, pdf_content: bytes) -> str:
        """
        Extract text from PDF document
        
        Args:
            pdf_content: PDF file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            pdf_file = io.BytesIO(pdf_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise
    
    def extract_text_from_docx(self, docx_content: bytes) -> str:
        """
        Extract text from DOCX document
        
        Args:
            docx_content: DOCX file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            docx_file = io.BytesIO(docx_content)
            doc = docx.Document(docx_file)
            text = ""
            
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise
    
    def extract_text_from_html(self, html_content: bytes) -> str:
        """
        Extract text from HTML document
        
        Args:
            html_content: HTML file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from HTML: {str(e)}")
            raise
    
    def extract_text(self, content: bytes, file_type: str) -> str:
        """
        Extract text based on file type
        
        Args:
            content: File content as bytes
            file_type: Type of file (pdf, docx, txt, html)
            
        Returns:
            Extracted text
        """
        if file_type == 'pdf':
            return self.extract_text_from_pdf(content)
        elif file_type == 'docx':
            return self.extract_text_from_docx(content)
        elif file_type == 'html' or file_type == 'htm':
            return self.extract_text_from_html(content)
        elif file_type == 'txt':
            return content.decode('utf-8')
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def create_fixed_size_chunks(
        self,
        text: str,
        chunk_size: int = 1000,
        overlap: int = 100
    ) -> List[str]:
        """
        Create fixed-size chunks with overlap
        
        Args:
            text: Text to chunk
            chunk_size: Maximum characters per chunk
            overlap: Number of characters to overlap
            
        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk.strip())
            start = end - overlap
        
        return [c for c in chunks if c]  # Remove empty chunks
    
    def create_semantic_chunks(
        self,
        text: str,
        max_chunk_size: int = 1000,
        overlap: int = 100
    ) -> List[str]:
        """
        Create chunks based on sentence boundaries (semantic chunking)
        
        Args:
            text: Text to chunk
            max_chunk_size: Maximum characters per chunk
            overlap: Number of words to overlap
            
        Returns:
            List of text chunks
        """
        # Split into sentences
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = ""
        previous_sentences = []
        
        for sentence in sentences:
            # Check if adding this sentence would exceed max size
            if len(current_chunk) + len(sentence) <= max_chunk_size:
                current_chunk += sentence + " "
                previous_sentences.append(sentence)
            else:
                # Save current chunk
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap
                if overlap > 0 and previous_sentences:
                    overlap_text = " ".join(previous_sentences[-overlap:])
                    current_chunk = overlap_text + " " + sentence + " "
                else:
                    current_chunk = sentence + " "
                
                previous_sentences = [sentence]
        
        # Add the last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def create_paragraph_chunks(self, text: str, max_chunk_size: int = 1500) -> List[str]:
        """
        Create chunks based on paragraph boundaries
        
        Args:
            text: Text to chunk
            max_chunk_size: Maximum characters per chunk
            
        Returns:
            List of text chunks
        """
        # Split into paragraphs
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            if len(current_chunk) + len(para) <= max_chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def create_sliding_window_chunks(
        self,
        text: str,
        window_size: int = 500,
        step_size: int = 250
    ) -> List[str]:
        """
        Create chunks using sliding window approach
        
        Args:
            text: Text to chunk
            window_size: Size of each window
            step_size: Step size for sliding
            
        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = min(start + window_size, text_length)
            chunk = text[start:end]
            chunks.append(chunk.strip())
            start += step_size
        
        return [c for c in chunks if c]
    
    def extract_metadata(
        self,
        text: str,
        file_name: str,
        s3_metadata: Optional[Dict] = None
    ) -> Dict:
        """
        Extract and generate metadata from document
        
        Args:
            text: Document text
            file_name: Original file name
            s3_metadata: S3 object metadata if available
            
        Returns:
            Dictionary of metadata
        """
        metadata = {
            'title': os.path.basename(file_name),
            'file_name': file_name,
            'document_length': len(text),
            'word_count': len(text.split()),
            'created_at': datetime.utcnow().isoformat()
        }
        
        # Add S3 metadata if available
        if s3_metadata:
            metadata['author'] = s3_metadata.get('author', 'Unknown')
            metadata['content_type'] = s3_metadata.get('ContentType', '')
            metadata['last_modified'] = s3_metadata.get('LastModified', '').isoformat() if s3_metadata.get('LastModified') else None
        
        # Calculate reading level (Flesch-Kincaid approximation)
        sentences = len(re.split(r'[.!?]+', text))
        words = len(text.split())
        syllables = sum(self._count_syllables(word) for word in text.split())
        
        if sentences > 0 and words > 0:
            reading_level = 0.39 * (words / sentences) + 11.8 * (syllables / words) - 15.59
            metadata['reading_level'] = round(reading_level, 2)
        
        return metadata
    
    def _count_syllables(self, word: str) -> int:
        """Simple syllable counter (approximation)"""
        word = word.lower()
        vowels = 'aeiou'
        count = 0
        previous_was_vowel = False
        
        for char in word:
            is_vowel = char in vowels
            if is_vowel and not previous_was_vowel:
                count += 1
            previous_was_vowel = is_vowel
        
        # Adjust for silent 'e'
        if word.endswith('e'):
            count -= 1
        
        # Every word has at least one syllable
        if count == 0:
            count = 1
        
        return count
    
    def generate_checksum(self, content: bytes) -> str:
        """
        Generate MD5 checksum for content
        
        Args:
            content: File content as bytes
            
        Returns:
            MD5 checksum string
        """
        return hashlib.md5(content).hexdigest()
    
    def process_document_from_s3(
        self,
        bucket: str,
        key: str,
        chunking_strategy: str = 'semantic'
    ) -> Dict:
        """
        Complete document processing pipeline from S3
        
        Args:
            bucket: S3 bucket name
            key: S3 object key
            chunking_strategy: Strategy for chunking (semantic, fixed, paragraph, sliding)
            
        Returns:
            Dictionary containing processed document information
        """
        # Download document
        response = self.s3.get_object(Bucket=bucket, Key=key)
        content = response['Body'].read()
        
        # Generate document ID and checksum
        document_id = str(uuid.uuid4())
        checksum = self.generate_checksum(content)
        
        # Determine file type
        file_extension = key.split('.')[-1].lower()
        
        # Extract text
        text = self.extract_text(content, file_extension)
        
        # Create chunks based on strategy
        if chunking_strategy == 'semantic':
            chunks = self.create_semantic_chunks(text)
        elif chunking_strategy == 'fixed':
            chunks = self.create_fixed_size_chunks(text)
        elif chunking_strategy == 'paragraph':
            chunks = self.create_paragraph_chunks(text)
        elif chunking_strategy == 'sliding':
            chunks = self.create_sliding_window_chunks(text)
        else:
            raise ValueError(f"Unknown chunking strategy: {chunking_strategy}")
        
        # Extract metadata
        s3_metadata = {
            'ContentType': response.get('ContentType', ''),
            'LastModified': response.get('LastModified'),
            'author': response.get('Metadata', {}).get('author', 'Unknown')
        }
        metadata = self.extract_metadata(text, key, s3_metadata)
        
        # Build result
        result = {
            'document_id': document_id,
            'bucket': bucket,
            'key': key,
            'checksum': checksum,
            'document_type': file_extension,
            'metadata': metadata,
            'chunks': chunks,
            'total_chunks': len(chunks)
        }
        
        logger.info(f"Processed document {key}: {len(chunks)} chunks")
        return result


if __name__ == "__main__":
    # Example usage
    processor = DocumentProcessor()
    
    # Example text
    sample_text = """
    Amazon Bedrock is a fully managed service that offers a choice of 
    high-performing foundation models (FMs) from leading AI companies like 
    AI21 Labs, Anthropic, Cohere, Meta, Stability AI, and Amazon via a single API.
    
    Along with a broad set of capabilities you need to build generative AI 
    applications with security, privacy, and responsible AI, Amazon Bedrock 
    helps you experiment with and evaluate top FMs for your use case.
    """
    
    chunks = processor.create_semantic_chunks(sample_text, max_chunk_size=200)
    print(f"Created {len(chunks)} chunks:")
    for i, chunk in enumerate(chunks):
        print(f"\nChunk {i+1}:")
        print(chunk[:100] + "...")




